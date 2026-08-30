#!/usr/bin/env python3
"""Build the Word deliverable for the co-authors from the v7 Methods content.

⚠ 정정 (2026-08-30): 옛 docstring 은 *"draft 가 정본이고 이 스크립트는 배치만 한다"*
라고 적었는데 **이 스크립트는 draft 를 읽지 않는다.**  본문·표가 아래 파이썬 상수에
따로 들어 있어서, 한쪽만 고치면 두 산출물이 조용히 갈라진다 (실제로 갈라져 있었다).

지금 규약은 이렇다:
  · 값의 정본은 `docs/reviews/table_s3_data_20260827.md` 다.
  · draft(`docs/manuscript/methods_simulation_v7_draft.md`)와 이 파일은 **둘 다 파생**이다.
  · `--selftest` 가 **draft 파일을 실제로 열어** 여기 적힌 핵심 수치가 거기에도 있는지
    대조한다 (§8).  갈라지면 빨간불이 난다 — 자기 신고가 아니라 실물 대조다.

Unresolved table entries are written as bracketed placeholders that say WHY they are
unresolved, so that a reader can tell "not yet measured" from "measured and small".

    python3 scripts/build_methods_docx.py [-o OUT.docx]
    python3 scripts/build_methods_docx.py --selftest
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

REPO = Path(__file__).resolve().parent.parent
DEFAULT_OUT = REPO / "docs" / "manuscript" / "Methods_simulation_v7_for_coauthors.docx"

# --- status --------------------------------------------------------------
STATUS = "PROVISIONAL — NOT FOR SUBMISSION"
STATUS_BODY = (
    "이 문서는 투고용 교체안이 아니라 쟁점을 결정하기 위한 검토 패키지입니다. "
    "적대 리뷰(R10)의 해제조건 8개 중 여섯이 해소됐고 둘이 부분입니다 — ⑤ Table S3 의 "
    "σ_ion 한 행, ⑦ 중 Figure 4a 그림. "
    "⚠ 다만 이 판정은 저희 자체 판정입니다. R11 을 올린 뒤 독립 재판정을 받지 "
    "못했으므로, ‘해소’ 는 ‘저희가 확인했다’ 이지 ‘제3자가 확인했다’ 가 아닙니다. "
    "Methods 본문(영문)은 그대로 붙여 넣을 수 있는 상태이고, 값이 없는 행은 대괄호로 "
    "사유를 적어 두었습니다."
)

RELEASE_CONDITIONS = [
    ("①", "두 PTFE 규약을 동등한 sensitivity 두 점으로 표기", "해소 (형식은 그 뒤 한 번 더 바뀜)",
     "굵은 글씨·reported·resolved 표기를 전부 제거했습니다. ⚠ 그 뒤 형식을 한 번 더 "
     "바꿨습니다 — 본문에서 두 값을 동등하게 병기하는 형식은 저희가 지어낸 것이고 출판 "
     "선례를 찾지 못했습니다. 같은 재료계·같은 코드의 선행 연구는 파라미터 민감도를 별도 "
     "절에 싣고 본문은 공칭값 하나를 씁니다. 그래서 본문은 공칭 규약 하나를 쓰고 같은 "
     "문단에 다른 값을 병기하며, 민감도는 Table S3c 로 옮겼습니다. ⚠ 사전등록된 검사는 "
     "그 규약을 ‘채택하지 않는다’ 로 판정했고 그 판정은 그대로입니다 — 본문이 그 값을 "
     "쓰는 것은 편집 결정이지 판정 번복이 아니며, 원고가 그렇게 적습니다."),
    ("②", "1.35 GPa 의 출처를 정확히 서술", "해소",
     "'경험적 접촉법칙 입력값이지 LPSCl 고유 물성이 아니다' 로 고쳤습니다."),
    ("③", "'하한(lower bound)' 표현 철회", "해소",
     "격자 미수렴은 '검토한 세분 구간에서 단조 증가' 로만 적고 전역 한계를 주장하지 않습니다."),
    ("④", "W4 32팔 원자료(JSON·receipt) 리포 커밋", "해소",
     "두 규약 각 16팔 + run_receipt + cohort_manifest + verdict_receipt 를 리포에 "
     "넣었습니다(docs/data/w4_ptfe_centerline_20260827, w4b_ptfe_off_20260827). 리포만으로 "
     "판정기를 돌려 비 1.307820 / 1.123672 를 재도출했고 파일 해시가 32/32 일치합니다. "
     "⚠ 다만 커밋된 것은 판정에 필요한 스칼라만 남긴 감사 패키지이지 원자료가 아닙니다 "
     "(팔당 131 MB → 5.8 kB, 필드 해는 버렸습니다). 솔버 재실행은 여전히 계산기의 원본이 "
     "필요하고, 매니페스트에 원본 해시만 남아 있습니다."),
    ("⑤", "σ_ion 과 Table S3 나머지 행", "부분 해소",
     "구조 지표 다섯 행 중 넷은 새 침대에서 측정해 표에 넣었습니다. σ_ion 한 행이 "
     "남았는데, 상태가 ‘미측정’ 이 아니라 ‘쟀는데 못 씁니다’ 입니다 — 아래 미결 항목 참조. "
     "면적 용량 행은 비용량이 실험 쪽 값이라 협업자 소관으로 넘겼습니다(저희가 낼 수 "
     "있는 면적 하중은 확정)."),
    ("⑥", "ε_union · thickness 의 연산 정의 명시", "해소",
     "두 값이 어떻게 계산되는지, 통상적인 전극 porosity 와 무엇이 다른지 적었습니다."),
    ("⑦", "그림·표 생성기 배선 및 Figure 4b 재작도", "부분 해소",
     "Figure 4b 는 두 규약을 병기하는 형태로 재작도했고 생성기를 리포에 넣었습니다. "
     "Figure 4a 는 캡션만 고쳤고 그림은 아직 교체하지 않았습니다 — 침대가 재압밀되어 "
     "입자 위치가 바뀌었으므로 같은 경로로 다시 뽑아야 합니다."),
    ("⑧", "두 판(설명형·압축형) 사실 일치 + 외부 DOCX 캡션 감사", "해소",
     "두 판의 사실 일치를 맞췄고, SI 캡션을 전수 감사해 전 과정을 ‘DEM’ 으로 부르는 "
     "오귀속 세 곳을 찾았습니다(Figure S16 설명, Table S3 제목, 본문 전송 문단). "
     "Figure S17·S18 에는 오귀속이 없었습니다."),
]

CHANGES = [
    ("1", "전 과정을 'DEM' 으로 부름",
     "DEM · MPM · voxel FV 를 이름으로 분리",
     "실제로 세 도구입니다. 표 제목과 그림 설명이 전부 DEM 이라 공저자도 MPM 이 쓰인 것을 "
     "볼 수 없었습니다."),
    ("2", "'E 를 24 → 1.35 GPa 로 연화' 가 문장 앞",
     "치밀화 표적이 앞, E 는 그 표적을 맞추는 노브로 뒤",
     "실제로 한 일이 그것입니다. 두 값은 Table S2 에 그대로 남깁니다 — 감추는 것이 아니라 "
     "순서를 바로잡는 것입니다."),
    ("3", "'paired mean with its standard error'",
     "origin-phase spread 와 range, 표준오차·신뢰구간은 명시적으로 부정",
     "8개 origin 은 한 침대의 완전한 2×2×2 factorial 이라 복제 오차 자유도가 0 입니다. "
     "sd/√8 은 표준오차가 아닙니다."),
    ("4", "PTFE 규약·탄소 전도도 규약이 암묵적",
     "둘 다 명시, PTFE 는 두 규약을 나란히",
     "이득의 크기가 규약에 크게 의존합니다(비 1.124 ↔ 1.308). 하나만 적으면 독자가 그것을 "
     "물리로 읽습니다."),
]

METHODS_FULL = [
    ("Microstructure reconstruction and transport simulation.", None),
    (None,
     "Three-dimensional SBE and DBE microstructures were built in two stages and then used as "
     "the geometry for a separate transport calculation, so that three distinct tools are "
     "involved: a discrete element method (DEM) for the rigid-particle packing, a material point "
     "method (MPM) for the plastic deformation of the electrolyte, and a finite-volume solver on "
     "a voxel grid for the effective conductivities."),
    ("Stage 1 — DEM packing.",
     "Rigid-sphere packing was computed in LIGGGHTS using 1,271 NCM811 spheres (radius 2.5 μm) "
     "and 146,420 LPSCl spheres (radius 0.5 μm), sized after the experimental powders and mixed "
     "at 70:27 by weight in a 50 × 50 μm² domain, compacted under displacement control. "
     "Rigid-sphere contacts do reproduce particle rearrangement, but not the plastic flattening, "
     "fracture and grain-boundary deformation that also densify sulfide powders, so the contact "
     "stiffness of LPSCl was selected against a densification target for sulfide cold pressing "
     "rather than taken from the dense material. The 1.35 GPa value is an empirical contact-law "
     "input, not an intrinsic LPSCl modulus and not an independent validation of the present "
     "SBE/DBE beds; the dense-material value (24 GPa) is listed beside it in Table S2. The ~10 % "
     "target is derived from composite and glass literature rather than measured directly on pure "
     "LPSCl at 300 MPa, and the 11–12 % contact overlap is a pure-SE simulation consistency "
     "result rather than a measured calibration target."),
    ("Stage 2 — MPM compaction.",
     "Plastic deformation was then resolved on the fixed DEM skeleton with a GPU-accelerated MPM. "
     "VGCF fibres, PTFE fibrils and SDCP particles were present in the material-point cloud "
     "during this stage at the experimental weight fractions, so their stiffness enters the "
     "compaction rather than being added afterwards. The deviatoric plasticity follows a J2 "
     "(von Mises) model with a yield strength of 0.30 GPa; the elastic pair E = 1.53 GPa and "
     "ν = 0.49 is a model choice that sets K = 25.5 GPa and G = 0.51 GPa, confining the softening "
     "to shear while leaving the bulk response at a dense-solid value. The resulting beds are more "
     "compacted than the experimental porosity anchor, so this parameterisation is not a "
     "validation of the present SBE/DBE geometry."),
    ("Stage 3 — voxel transport.",
     "Each microstructure was rasterized onto a cubic grid with a voxel edge of 0.15 μm. "
     "Adjacent conducting voxels were coupled through harmonic-mean conductances and the potential "
     "field obtained from ∇·(σ∇φ) = 0, with 1 V applied between the separator (φ = 0) and "
     "current-collector (φ = 1 V) faces and the remaining boundaries insulating; the effective "
     "conductivity was taken from the total current. NCM811, VGCF and SDCP carried the electronic "
     "network and LPSCl and SDCP the ionic network. The insulating binder was treated under two "
     "conventions, reported as equivalent sensitivity points rather than one primary result: "
     "omitted from the electronic grid, and with its centerline voxels excluded from conduction. "
     "Neither is established as closer to a real thin coating — the one-cell centerline "
     "under-represents the coating’s spatial extent while over-blocking where it is stamped, and "
     "the diameter-aware variant is not implemented. Note that omitting the binder from the "
     "conduction grid does not remove it from the model: its mass and stiffness are present in the "
     "DEM–MPM bed, and only its direct insulating exclusion on the electronic grid is absent."),
    ("Conductivity of the carbon network.",
     "The conductivity assigned to VGCF is an effective network value, not a fibre material "
     "constant. Voxelisation fuses touching fibres into shared cells and therefore removes the "
     "fibre–fibre contact resistance that dominates a real carbon network — the two orders of "
     "magnitude between the compressed-powder (≈ 83 S cm⁻¹) and single-filament "
     "(≈ 10⁴ S cm⁻¹) values of VGCF-H is essentially that contact resistance. The powder-scale "
     "value (100 S cm⁻¹) was therefore adopted and rescaled to the voxel grid so that the axial "
     "conductance of a fibre is preserved (78.5 S cm⁻¹ at 0.15 μm)."),
    ("Grid-origin ensemble and reported statistics.",
     "Because the voxel grid samples the same microstructure differently depending on where its "
     "origin falls, each electrode was solved at all eight half-voxel origin shifts of a "
     "2 × 2 × 2 factorial, the SBE and DBE sharing the same origins so that ratios are formed "
     "pair by pair. The eight phases are a complete factorial of a single bed rather than "
     "independent replicates, so ratios are reported as the mean over the eight prescribed phases "
     "together with the spread across them and the observed range; no standard error or confidence "
     "interval is implied. All arms reached the solver convergence criterion."),
    ("Values under the two binder conventions.",
     "With PTFE centerline voxels excluded (exact-zero sensitivity convention) the effective "
     "electronic conductivity is 54.0 mS cm⁻¹ (SBE) and 70.6 mS cm⁻¹ (DBE), a paired ratio of "
     "1.308 (spread 0.003 across the eight origin phases; observed range 1.302–1.310). Leaving "
     "PTFE unresolved gives 72.3 and 81.3 mS cm⁻¹, a ratio of 1.124 (spread 0.003; range "
     "1.120–1.127). The two differ only in whether the binder’s centerline voxels are excluded "
     "from conduction — a machine-checked contract confirms that no other parameter differs — and "
     "neither is established as closer to a real thin coating: the one-cell centerline "
     "under-represents the coating’s spatial extent while over-blocking where it is stamped. They "
     "are therefore reported as two equivalent model-form sensitivity points; the direction of the "
     "change is common to both, its magnitude is not. Where the main text needs a single number it "
     "quotes the centerline convention; that is a reporting choice and not a determination that "
     "either convention is closer to a real coating. Ohmic loss per phase was evaluated as "
     "Σ gₖ Δφₖ², summed over the voxel-to-voxel connections belonging to that phase."),
    ("Limitations.",
     "The absolute conductivities have not been calibrated against a composition-matched "
     "measurement and should be read as the output of an idealised bulk model: the solver places "
     "no contact resistance at any interface — between active particles, between active material "
     "and carbon, or at the current collector — so it does not reproduce the quantity a "
     "two-terminal DC-polarisation measurement returns. The ratio is not grid-converged: refining "
     "the voxel edge increases it monotonically without following a power law, so the reported "
     "gain is larger at finer voxels over the refinement interval examined; neither a continuum "
     "extrapolation nor a global bound is established. Refinement here changes two things at "
     "once, since the diameter-preserving carbon conductivity is itself a function of the voxel "
     "edge; a separate arm that moved the carbon conductivity alone over the same factor shifts "
     "the ratio in the opposite direction and by a small fraction of the observed change, so the "
     "geometric part of the grid dependence is if anything larger than the combined figure. Explicitly restoring the additive contacts "
     "that voxelisation drops recovers only about a fifth of that grid dependence — measured under "
     "the binder-omitted convention only, and not transferable to the exact-zero convention. The "
     "magnitude is also conditional on the carbon conductivity being treated as an effective "
     "network constant: at the single-filament value, a hundredfold higher, the ordering reverses. "
     "That upper arm is not better physics but a doubly idealised sensitivity at one origin phase, "
     "assuming perfect fibre–fibre contact on top of equipotential fibres; a single scalar can "
     "absorb part of the missing resistance numerically but is not identified as a fibre–fibre "
     "contact parameter. Two further limits apply: the beds are more compacted than the "
     "experimental porosity anchor, and the specimen provenance of the SDCP conductivity "
     "(250 S cm⁻¹) is not established."),
]

METHODS_COMPACT = [
    ("Microstructure reconstruction and transport simulation.",
     "Three-dimensional SBE and DBE microstructures were built with a discrete element method "
     "(DEM) for the rigid-particle packing and a material point method (MPM) for the plastic "
     "deformation of the electrolyte, and the effective conductivities were then obtained with a "
     "finite-volume solver on a voxel grid. Rigid-sphere packing was computed in LIGGGHTS from "
     "1,271 NCM811 spheres (r = 2.5 μm) and 146,420 LPSCl spheres (r = 0.5 μm), mixed 70:27 by "
     "weight in a 50 × 50 μm² domain and compacted under displacement control. Rigid spheres do "
     "rearrange but cannot flatten, fracture or deform grain boundaries, so the LPSCl contact "
     "stiffness was selected against a densification target for sulfide cold pressing rather than "
     "taken from the dense material: 1.35 GPa is an empirical contact-law input, not an intrinsic "
     "modulus and not a validation of these beds, and 24 GPa is listed alongside it in Table S2. "
     "The ~10 % target is derived from composite and glass literature rather than measured on pure "
     "LPSCl, and the 11–12 % overlap is a pure-SE simulation result, not a measured target. "
     "Plastic deformation was then resolved on the fixed DEM skeleton by MPM with J2 plasticity "
     "(σ_y = 0.30 GPa); the elastic pair E = 1.53 GPa, ν = 0.49 is a model choice giving "
     "K = 25.5 GPa and G = 0.51 GPa, confining the softening to shear. VGCF, PTFE and SDCP were "
     "present in the material-point cloud during compaction at the experimental weight fractions. "
     "The resulting beds are more compacted than the experimental porosity anchor."),
    (None,
     "Each microstructure was rasterized at a voxel edge of 0.15 μm; adjacent conducting voxels "
     "were coupled through harmonic-mean conductances and ∇·(σ∇φ) = 0 solved with 1 V between "
     "the separator and current-collector faces, the remaining boundaries insulating. NCM811, VGCF "
     "and SDCP carried the electronic network and LPSCl and SDCP the ionic network. The insulating "
     "binder was treated under two conventions reported as equivalent sensitivity points — omitted "
     "from the electronic grid, and with its centerline voxels excluded — since neither is "
     "established as closer to a real thin coating: the one-cell centerline under-represents the "
     "coating’s extent while over-blocking where stamped, and omitting it removes only the "
     "electronic exclusion, the binder’s mass and stiffness remaining in the bed. The VGCF "
     "conductivity is an effective network value rather than a fibre constant, since voxelisation "
     "removes the fibre–fibre contact resistance that separates the powder (≈ 83 S cm⁻¹) and "
     "single-filament (≈ 10⁴ S cm⁻¹) values; the powder-scale value was adopted and rescaled to "
     "preserve axial fibre conductance (78.5 S cm⁻¹)."),
    (None,
     "Each electrode was solved at all eight half-voxel grid-origin shifts of a 2 × 2 × 2 "
     "factorial, SBE and DBE sharing the same origins so that ratios are paired. These eight "
     "phases are a complete factorial of a single bed rather than independent replicates, so ratios "
     "are given as the mean over the prescribed phases with the spread and observed range; no "
     "standard error and no confidence interval are implied. Ohmic loss per phase was evaluated as "
     "Σ gₖ Δφₖ². The two binder conventions give 72.3/81.3 mS cm⁻¹ (ratio 1.124) and "
     "54.0/70.6 mS cm⁻¹ (ratio 1.308): the direction is common to both, the magnitude is not. "
     "Absolute conductivities are those of an idealised bulk model with no interfacial contact "
     "resistance anywhere and are not composition-matched to a measurement. The ratio is not "
     "grid-converged and grew at finer voxels over the refinement interval examined; no continuum "
     "extrapolation or global bound is established. Refinement moves the voxel edge and the "
     "diameter-preserving carbon conductivity together; an arm isolating the latter shifts the "
     "ratio the other way, so the geometric part is not smaller than the combined figure. Restoring the additive contacts that "
     "voxelisation drops recovers about a fifth of that dependence, measured under the "
     "binder-omitted convention only. The magnitude is also conditional on the carbon conductivity "
     "being an effective network constant: at the single-filament value the ordering reverses, but "
     "that arm is not better physics — it is a doubly idealised sensitivity at one origin phase. "
     "Two further limits: the beds are more compacted than the experimental porosity anchor, and "
     "the specimen provenance of the SDCP conductivity (250 S cm⁻¹) is not established."),
]

# --- tables ---------------------------------------------------------------
PENDING_ION = "[ 보류 — 쟀으나 입력 규약이 두 침대를 비대칭으로 만든다 (아래 미결 참조) ]"
PENDING_CAP = "[ 협업자 소관 — 면적 하중 0.015904 g cm⁻² 확정, 비용량은 원고 밖 ]"

TABLE_S2 = [
    ("Block", "Rows carried in that block", "Role label"),
    ("DEM (packing)",
     "Domain 50 × 50 μm² · NCM811 r = 2.5 μm, E = 140 GPa · "
     "LPSCl r = 0.5 μm, E(dense) = 24 GPa, E(DEM contact) = 1.35 GPa",
     "E(DEM contact): Empirical contact-law input"),
    ("MPM (plastic compaction)",
     "E = 1.53 GPa · ν = 0.49 (these give K = 25.5 GPa, G = 0.51 GPa) · σ_y = 0.30 GPa",
     "E, ν: Model choice · σ_y: Selected against densification target"),
    ("Voxel transport",
     "Voxel edge 0.15 μm · σ_e(NCM811) 1.0 × 10⁻² · σ_ion(LPSCl) 3.0 × 10⁻³ · "
     "σ_e(VGCF, powder) 1.0 × 10² · σ_e(VGCF, voxel diameter-preserving) 78.5 · "
     "σ_e(SDCP) 250 · PTFE 0  [S cm⁻¹]",
     "σ_e(VGCF): Effective network constant, not a fibre material constant"),
]

TABLE_S3 = [
    ("Parameter", "SBE", "DBE", "Unit"),
    ("Thickness", "72.53", "72.53", "μm"),
    ("ε_union (simulation-geometry diagnostic)", "7.86", "7.37", "%"),
    ("σ_ele,eff — PTFE omitted from the electronic grid (legacy/default convention)",
     "72.3", "81.3", "mS cm⁻¹"),
    ("σ_ele,eff — PTFE centerline voxels excluded (exact-zero sensitivity convention)",
     "54.0", "70.6", "mS cm⁻¹"),
    ("σ_ele ratio, paired over 8 origin phases (omitted · centerline-excluded)",
     "1.124", "1.308", "—"),
    ("└ spread / observed range", "0.003 / 1.120–1.127", "0.003 / 1.302–1.310", "—"),
    ("σ_ion,eff", PENDING_ION, PENDING_ION, "mS cm⁻¹"),
    ("SE coverage of AM (Tabor band)", "86.6", "86.6", "%"),
    ("VGCF coverage of AM", "13.1", "15.5", "%"),
    ("Median CBD contacts per AM particle — conductive phases only (VGCF + SDCP)",
     "74", "86", "ea"),
    ("└ same count with the insulating binder included in the domain", "80", "88", "ea"),
    ("Electronic connectivity", "100", "100", "%"),
    ("Areal capacity", PENDING_CAP, PENDING_CAP, "mAh cm⁻²"),
]

TABLE_S3_FOOT = (
    "Coverage is reported in the Tabor band; the three coverage conventions used elsewhere in "
    "this project measure different quantities and must not be mixed. The contact count depends "
    "on whether the insulating binder is counted as part of the conductive-binder domain, so "
    "both counts are given: the increase from SBE to DBE is +16 % counting only the conductive "
    "phases and +10 % counting the binder as well. Neither reproduces the absolute numbers in "
    "v6 (433 and 517); what those counted is not recorded, and the direction and relative size "
    "of the change do agree. The bed itself reports a quasi-static violation — the platen moves "
    "at a substantial fraction of the compressional wave speed — so thickness and ε_union are "
    "absolute values that need re-measurement, while the ratio is common-mode and unaffected."
)

TABLE_S3B = [
    ("PTFE convention", "σ_ele SBE", "σ_ele DBE", "Ratio", "Spread", "Range"),
    ("PTFE omitted from the electronic grid (legacy/default convention)",
     "72.3", "81.3", "1.124", "0.003", "1.120–1.127"),
    ("PTFE centerline voxels excluded (exact-zero sensitivity convention)",
     "54.0", "70.6", "1.308", "0.003", "1.302–1.310"),
]

TABLE_S3B_FOOT = (
    "Both settings were evaluated on the same beds with the same code and grid, differing only in "
    "whether the insulating binder occupies conduction cells; a machine-checked contract confirms "
    "that no other parameter differs. The setting used in the main text stamps a one-voxel "
    "centerline through each binder fibril and removes those cells from conduction, which accounts "
    "for 43 % of the binder volume present in the bed; omitting the binder accounts for none of it. "
    "It was selected for reporting but is not calibrated: a pre-registered test asked whether "
    "either setting reproduces the binder volume within a factor of two and neither does, so the "
    "quantity that would decide which setting is closer to the real film has not been measured — "
    "blocking follows connectivity, not volume. The increase from SBE to DBE is obtained under "
    "both settings; its magnitude is not. The binder's mass and stiffness are present in the "
    "DEM–MPM bed under either, since the setting affects only the conduction grid."
)

MAINTEXT_EDITS = [
    ("Location", "v6 as written", "Proposed"),
    ("Transport paragraph",
     "“… the simulated effective σ_ele increases from 1.98 to 3.00 S cm⁻¹”",
     "One nominal setting in the body, with the alternative in the same paragraph: “the simulated "
     "effective σ_ele increases from 54.0 to 70.6 mS cm⁻¹ (a ratio of 1.308) with the binder’s "
     "centerline voxels excluded from conduction; omitting the binder from the electronic grid "
     "instead gives 72.3 → 81.3 mS cm⁻¹ (1.124). The setting used here was selected for reporting "
     "but is not calibrated (Table S3c).” — ⚠ 형식 근거: 같은 재료계·같은 코드의 선행 연구가 "
     "파라미터 민감도를 별도 절에 싣고 본문은 공칭값 하나를 씁니다. 두 값을 본문에서 "
     "동등하게 병기하는 형식은 저희가 찾은 선례가 없습니다."),
    ("Transport paragraph",
     "“… σ_ion … 0.203 and 0.215 mS cm⁻¹”",
     "[ 보류 — 이온 전용 런 완료 전까지 수치를 넣지 않음 ]"),
    ("Transport paragraph",
     "“reconstructed using a discrete element method (DEM)”",
     "“generated by DEM packing and MPM compaction” — ‘reconstructed’ implies tomographic "
     "reconstruction of a real specimen, which this is not."),
    ("Figure 4a caption", "“DEM-reconstructed SBE and DBE geometries”",
     "“SBE and DBE composite-cathode geometries generated by DEM particle packing and compacted "
     "by MPM” — 두 군데를 고칩니다: ‘reconstructed’ 는 실물 단층촬영 재구성을 암시하고, "
     "‘DEM’ 은 세 도구 중 하나만 부릅니다."),
    ("Figure 4a 그림", "출처가 기록돼 있지 않던 그림",
     "[ 교체 필요 — 침대가 재압밀되어 입자 위치가 바뀌었습니다. 지금 그림이 옛 침대라면 "
     "Table S3 이 보고하는 것과 다른 미세구조입니다 ]"),
    ("Figure 4b", "Plot of the superseded σ values",
     "재작도 완료 — 두 규약을 병기하는 형태로 다시 그렸고 생성기를 리포에 넣었습니다."),
    ("Table S3 title", "“… obtained from the DEM simulations”",
     "“Structural metrics from the DEM–MPM geometry and transport metrics from the voxel "
     "finite-volume solver”"),
    ("Figure S16 caption", "“Reconstructed SBE and DBE geometries used for the DEM simulations”",
     "“SBE and DBE geometries generated by DEM packing and MPM compaction” — 두 군데 "
     "(‘Reconstructed’ + 전 과정을 DEM 으로 부름)."),
    ("Figure S17 · S18 captions", "“Simulated ionic / electronic current-density distributions”",
     "오귀속 없음. 솔버 이름을 붙이는 편이 낫습니다(선택): “… computed with the voxel "
     "finite-volume solver”."),
]

OPEN_ITEMS = [
    ("σ_ion (Table S3) — ‘미측정’ 이 아니라 ‘쟀는데 못 씁니다’",
     "같은 침대·같은 origin 으로 이온까지 푸는 런을 돌렸고 값이 나왔습니다. 그런데 비가 "
     "1 보다 작게 나오는데, 그것이 물리가 아니라 입력 규약의 비대칭입니다. 생산 규약에서 "
     "바인더는 전도 격자에 한 셀도 들어가지 않습니다 — 그러면 바인더를 두 배로 가진 SBE 쪽 "
     "절연 부피를 더 많이 봐주게 됩니다. 동시에 SDCP 는 DBE 에만 있으면서 자기 자리의 "
     "전해질을 이온 전도도가 3 분의 1 인 상으로 바꿉니다. 두 힘이 같은 방향으로 비를 "
     "끌어내립니다. ⚠ 그러므로 남은 팔이 끝나도 값은 바뀌지 않습니다 — 팔 수가 아니라 "
     "입력이 정하는 값입니다. 채우면 ‘SDCP 가 이온 전도를 떨어뜨린다’ 를 싣게 되는데, "
     "그것은 모델이 바인더를 안 그린 결과입니다. 실물에서 SDCP 는 절연 바인더를 대체하고 "
     "모델에서는 전해질을 대체합니다 — 치환 상대가 다릅니다. "
     "★ 문헌이 방향을 확인해 줍니다: 같은 재료계에서 PTFE 1 wt% 가 이온 전도도를 26 % "
     "떨어뜨린다는 실측이 있고, 그 두 로딩이 정확히 저희 SBE·DBE 의 바인더 함량입니다. "
     "그것만으로도 비가 1 보다 커야 합니다 — 시뮬레이션과 부호가 반대입니다. "
     "해제는 펠릿 보정으로 σ_ion(SDCP) 와 바인더 차단 길이를 앵커한 뒤 다시 도는 것입니다."),
    ("Areal capacity — v6 의 두 값은 이 침대에서 성립하지 않습니다 (협업자 소관)",
     "SI v6 은 두 전극에 서로 다른 면적 용량을 적었는데, 두 전극이 같은 AM 골격과 같은 "
     "두께를 공유하므로 성립할 수 없습니다. ⚠ 이 사실은 누가 값을 채우든 그대로입니다. "
     "저희가 낼 수 있는 것은 면적 하중까지이고 0.015904 g cm⁻² 로 확정했습니다(두 전극 "
     "동일). 비용량은 시뮬레이션 산출물이 아니라 실험 쪽 값이라 이 행은 협업자가 "
     "채웁니다 — 시뮬레이션 축의 미결이 아닙니다."),
    ("Figure 4a 그림 (해제조건 ⑦ 잔여)",
     "캡션은 고쳤지만 그림은 그대로입니다. 침대를 재압밀하면서 첨가제 탄성계수 세대가 "
     "바뀌었고, 벌크 지표는 거의 안 움직였지만 입자 위치는 바뀌었습니다. 지금 그림이 옛 "
     "침대에서 나온 것이라면 Table S3 이 보고하는 것과 다른 미세구조를 보여주게 됩니다."),
    ("AM 입자당 CBD 접촉 수 — 규약이 이득까지 바꿉니다",
     "본문이 이 수를 기전 근거로 직접 인용합니다. 절연 바인더를 ‘conductive binder domain’ "
     "에 넣느냐에 따라 증가율이 16 % 와 10 % 로 갈립니다. 규약을 함께 적지 않으면 인용할 "
     "수 없습니다. 또 v6 의 절대값(433 · 517)은 재현되지 않습니다 — v6 이 무엇을 셌는지 "
     "기록이 없습니다. 방향과 상대 크기는 정합합니다."),
    ("격자 수렴",
     "이득이 격자를 조일수록 단조 증가하고 멈추지 않습니다. 증분 비가 어떤 멱법칙과도 맞지 "
     "않아 외삽이 성립하지 않으므로, 검토한 세분 구간에서 단조 증가한다고만 적고 전역 한계를 "
     "주장하지 않습니다."),
    ("탄소 전도도 가정의 여유",
     "현재 결론은 σ(VGCF) 를 유효 망 상수로 볼 때 성립합니다. 단섬유 값을 넣으면 이득의 "
     "부호가 뒤집히므로 — 망이 이미 완전하면 SDCP 가 보탤 것이 없고 오히려 부피를 차지합니다 "
     "— 이 가정은 결론에 실질적입니다. Limitations 에 한 문단으로 들어가 있고, 감도 공정을 "
     "런 전에 등록해 두었습니다."),
    ("준정적 위반",
     "침대 자체가 준정적 조건 위반을 보고합니다. 두 전극이 같은 속도로 눌렸으므로 비는 "
     "공통 모드로 상쇄되지만, 두께와 ε_union 은 절대값이라 영향을 받습니다. Limitations 에 "
     "적었습니다."),
    ("독립 재판정",
     "R11 을 올린 뒤 적대 리뷰를 더 받지 못했습니다. 위 해제조건의 ‘해소’ 는 전부 저희 "
     "자체 판정입니다."),
]


# --- rendering ------------------------------------------------------------
def _shade(cell, hexcolor):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _add_table(doc, rows, widths=None, small=False):
    from docx.shared import Pt
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, text in enumerate(row):
            cells[j].text = ""
            p = cells[j].paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(8 if small else 9)
            if i == 0:
                run.bold = True
            if str(text).startswith("["):
                run.italic = True
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
        if i == 0:
            for c in cells:
                _shade(c, "E8E8E8")
    if widths:
        from docx.shared import Cm
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Cm(w)
    return t


def build(out_path: Path) -> Path:
    import docx
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2.0)
        section.left_margin = section.right_margin = Cm(2.0)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)

    def h(text, level=1):
        p = doc.add_heading(text, level=level)
        for r in p.runs:
            r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        return p

    def para(text, *, italic=False, bold=False, size=10, indent=0.0, space=6):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.italic, r.bold = italic, bold
        r.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(space)
        if indent:
            p.paragraph_format.left_indent = Cm(indent)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p

    # --- title
    t = doc.add_paragraph()
    r = t.add_run("Methods (simulation) — revision v7")
    r.bold = True
    r.font.size = Pt(17)
    doc.add_paragraph("SBE / DBE microstructure reconstruction and transport simulation — "
                      "replacement text and revised SI tables for co-author review, 2026-08-30")

    # --- status banner
    p = doc.add_paragraph()
    r = p.add_run(STATUS)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
    para(STATUS_BODY, size=9.5)

    h("해제조건 8개 — 현재 상태", 2)
    _add_table(doc, [("#", "조건", "상태", "설명")] + [
        (n, cond, state, why) for n, cond, state, why in RELEASE_CONDITIONS
    ], widths=[0.9, 5.0, 1.8, 9.0], small=True)

    doc.add_page_break()

    # --- what changed
    h("1. v6 대비 무엇이 바뀌나 (넷)", 1)
    _add_table(doc, [("#", "v6", "v7", "왜 바꾸는가")] + list(CHANGES),
               widths=[0.9, 4.2, 4.6, 7.0], small=True)

    para("PTFE 규약에 대해서는 한 가지를 분명히 해 둡니다. 초안은 처음에 '안 그림' 을 주 규약으로 "
         "썼는데 이유가 편집 편의였고, 지적을 받고 '차단' 으로 뒤집었는데 이번에는 두 값을 본 뒤에 "
         "큰 쪽으로 옮긴 것이라 결과 독립이 아니었습니다. 그래서 지금은 어느 쪽도 주 규약으로 "
         "지정하지 않고 동등한 두 sensitivity 점으로 적습니다. 더 근본적으로, centerline 규약은 "
         "PTFE 를 제대로 해상한 규약이 아닙니다 — 한 셀 폭 중심선을 찍고 그 셀을 정확히 0 으로 "
         "제거하는 방식이라, 얇은 코팅의 공간 범위는 과소 표현하면서 찍힌 셀에서는 차단을 과대 "
         "표현합니다. 직경을 인식하는 변형은 아직 구현돼 있지 않습니다.", size=9.5)

    doc.add_page_break()

    # --- Methods full
    h("2. Methods — 설명형 (Full)", 1)
    para("독자가 이해할 수 있도록 단계를 나눠 쓴 판입니다. 지면이 허용되면 이쪽을 씁니다.",
         italic=True, size=9)
    for lead, body in METHODS_FULL:
        p = doc.add_paragraph()
        if lead:
            r = p.add_run(lead + " ")
            r.bold = True
            r.font.size = Pt(10)
        if body:
            r = p.add_run(body)
            r.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_page_break()

    # --- Methods compact
    h("3. Methods — 압축형 (Compact)", 1)
    para("같은 사실을 담되 문장 수만 줄인 판입니다. 한정어는 하나도 빼지 않았습니다 — "
         "압축은 문장을 줄이는 것이지 유보를 빼는 것이 아닙니다.", italic=True, size=9)
    for lead, body in METHODS_COMPACT:
        p = doc.add_paragraph()
        if lead:
            r = p.add_run(lead + " ")
            r.bold = True
            r.font.size = Pt(10)
        r = p.add_run(body)
        r.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_page_break()

    # --- Table S2
    h("4. Table S2 — 제목과 블록 분리", 1)
    para("제목: “Material parameters used for the DEM simulations” → "
         "“Material and numerical parameters used for the DEM–MPM–voxel transport workflow”.",
         size=9.5)
    _add_table(doc, TABLE_S2, widths=[3.2, 8.0, 5.5], small=True)
    para("E(dense) 24 GPa 와 E(DEM contact) 1.35 GPa 는 둘 다 남깁니다. VGCF 의 두 행(powder / "
         "voxel diameter-preserving)도 그대로 둡니다 — 규약이 표에서 보여야 합니다. 그리고 세 "
         "값의 역할이 서로 다르므로 'Calibrated' 한 라벨로 뭉치지 않습니다: 접촉 E 는 경험적 "
         "접촉법칙 입력값, MPM 의 E·ν 는 모델 선택, σ_y 는 치밀화 표적에 맞춰 고른 값입니다.",
         size=9)

    # --- Table S3
    h("5. Table S3 — σ_ele · 구조 지표 갱신, 두 행 보류", 1)
    para("제목: “… obtained from the DEM simulations” → “Structural metrics from the DEM–MPM "
         "geometry and transport metrics from the voxel finite-volume solver”.", size=9.5)
    _add_table(doc, TABLE_S3, widths=[8.2, 4.2, 4.2, 2.1], small=True)
    para("각주: " + TABLE_S3_FOOT, italic=True, size=8.5)
    para("대괄호로 남긴 칸은 값이 없는 칸입니다 — 작아서 생략한 것이 아니라 아직 재지 않았습니다. "
         "옛 값을 그대로 옮기지 않은 이유는, 그렇게 하면 한 표 안에 두 세대의 침대가 섞이기 "
         "때문입니다. σ_ele 만 새 침대이고 나머지가 옛 침대이면 독자는 그것을 구별할 수 없습니다.",
         size=9)
    para("ε_union 과 thickness 의 연산 정의:", bold=True, size=9.5, space=2)
    for line in [
        "· ε_union = 1 − V_solid / (A · (z_plate − z_floor)). 분자는 겹침을 한 번만 세는 합집합 "
        "부피이고 분모는 바닥판과 플래튼 사이의 상자 부피입니다. 통상적인 전극 porosity 가 "
        "아니며, 실험 앵커(~15.6 %) 대비 과압축입니다.",
        "· Thickness 는 운동학적 정지 규칙에서 플래튼이 멈춘 위치로 정해지는 값입니다. 응력 "
        "평형에서 창발한 두께가 아닙니다.",
        "· 두 침대의 thickness 가 같다는 것이 과압축이 비에서 상쇄된다는 뜻은 아닙니다. 같은 "
        "속도·같은 정지 위치는 like-for-like 입력을 보장할 뿐입니다.",
    ]:
        para(line, size=9, indent=0.5, space=2)

    # --- Table S3b
    h("6. Table S3c (신설) — 바인더 표현 민감도", 1)
    _add_table(doc, TABLE_S3B, widths=[6.6, 2.2, 2.2, 1.8, 1.8, 2.6], small=True)
    para("각주: " + TABLE_S3B_FOOT, italic=True, size=8.5)

    doc.add_page_break()

    # --- main text edits
    h("7. 본문에서 함께 고쳐야 할 곳", 1)
    para("아래는 협업자 문서의 문장이라 저희가 바꾸지 않고 지목만 합니다.", italic=True, size=9)
    _add_table(doc, MAINTEXT_EDITS, widths=[3.4, 5.6, 7.7], small=True)

    # --- open items
    h("8. 아직 해결되지 않은 것 — 전체 설명", 1)
    for title, body in OPEN_ITEMS:
        p = doc.add_paragraph()
        r = p.add_run(title + " — ")
        r.bold = True
        r.font.size = Pt(9.5)
        r = p.add_run(body)
        r.font.size = Pt(9.5)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    h("9. 수치의 출처", 2)
    para("이 문서의 σ_ele 값은 2026-08-27 침대에서 vox 0.15 μm, SDCP 를 참 직경 구로 스탬프한 "
         "규약으로, 두 PTFE 규약 각각 8개 origin 씩 모두 16팔을 완주해 얻은 것입니다. 두 규약의 "
         "짝은 같은 침대·같은 코드이며 등록 축 외의 인자가 다르지 않다는 것을 기계 검사로 "
         "확인했습니다. 수치 정본은 리포의 docs/reviews/table_s3_data_20260827.md 이고, 이 "
         "문서는 그것을 옮겨 적은 것입니다 — 두 곳이 어긋나면 정본이 이깁니다.", size=9.5)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


# --- selftest -------------------------------------------------------------
def selftest() -> int:
    fails = []

    def chk(name, cond, detail=""):
        if cond:
            print(f"  ok   {name}")
        else:
            fails.append(name)
            print(f"  FAIL {name} {detail}")

    print("build_methods_docx selftest")

    # 1. every table is rectangular
    for label, tbl in [("S2", TABLE_S2), ("S3", TABLE_S3), ("S3b", TABLE_S3B),
                       ("maintext", MAINTEXT_EDITS)]:
        widths = {len(r) for r in tbl}
        chk(f"table {label} rectangular", len(widths) == 1, f"got {widths}")

    # 2. no retracted value leaks in (the ban registry patterns that touch this axis)
    banned = ["1.1232", "1.123191", "+12.3 %", "+12.32", "+52.0", "+42.15",
              "1.143817", "1.155448", "1.98", "3.00 S cm", "f_artifact", "×35.79"]
    blob = "\n".join([
        STATUS_BODY, TABLE_S3B_FOOT,
        *[c for row in RELEASE_CONDITIONS for c in row],
        *[c for row in CHANGES for c in row],
        *[str(c) for row in TABLE_S2 + TABLE_S3 + TABLE_S3B for c in row],
        *[b or "" for _, b in METHODS_FULL],
        *[a or "" for a, _ in METHODS_FULL],
        *[b for _, b in METHODS_COMPACT],
        *[c for _, c in OPEN_ITEMS],
    ])
    # main-text edit rows quote v6 deliberately; check them separately
    for pat in banned:
        chk(f"no retracted pattern {pat!r} in body", pat not in blob)

    # 3. the two Methods versions must agree on the numbers they state
    import re
    def nums(text):
        return set(re.findall(r"\d+\.\d+", text))
    full = " ".join(b or "" for _, b in METHODS_FULL)
    comp = " ".join(b for _, b in METHODS_COMPACT)
    for key in ["72.3", "81.3", "1.124", "54.0", "70.6", "1.308", "0.15", "1.35",
                "1.53", "0.49", "25.5", "0.51", "0.30", "78.5", "2.5", "0.5"]:
        chk(f"both versions state {key}", key in nums(full) and key in nums(comp))

    # 4. every unresolved cell must be a bracketed explanation, never blank or a number
    for row in TABLE_S3[1:]:
        for cell in row[1:3]:
            if cell.startswith("["):
                chk("pending cell explains itself",
                    "—" in cell and len(cell) > 10, cell)

    # 5. the structural rows that the main text cites must still be present
    for needed in ["CBD contacts per AM particle", "SE coverage of AM",
                   "VGCF coverage of AM", "Electronic connectivity", "Areal capacity"]:
        chk(f"row kept: {needed}", any(needed in r[0] for r in TABLE_S3))

    # 6. no standard-error language survives
    for bad in ["standard error", "confidence interval"]:
        occurrences = [t for t in (full, comp) if bad in t]
        chk(f"{bad!r} only ever negated",
            all("no standard error" in t or "no confidence interval" in t
                for t in occurrences))

    #  6b. 렌더되는 문자열에 마크다운 강조가 남으면 Word 에 별표가 그대로 찍힌다
    #      (para() 는 마크다운을 파싱하지 않는다).  실제로 한 번 찍혔다.
    _rendered = [STATUS_BODY, TABLE_S3_FOOT, TABLE_S3B_FOOT,
                 *[c for row in RELEASE_CONDITIONS for c in row],
                 *[c for _, c in OPEN_ITEMS],
                 *[str(c) for row in TABLE_S2 + TABLE_S3 + TABLE_S3B for c in row]]
    _md = [s[:60] for s in _rendered if "**" in s]
    chk("렌더 문자열에 마크다운 강조가 없다", not _md, _md[:2])

    # 7. release conditions must be exactly the eight
    chk("eight release conditions", len(RELEASE_CONDITIONS) == 8)

    #  8. ★★ draft 실물 대조 (2026-08-30).  옛 docstring 은 "draft 가 정본" 이라고 적으면서
    #     이 파일을 한 번도 열지 않았고, 그래서 두 산출물이 조용히 갈라졌다.  여기서는
    #     **파일을 실제로 열어** 이 표의 핵심 수치가 거기에도 있는지 본다.  자기 신고가
    #     아니라 실물 대조다 — 한쪽만 고치면 빨간불이 난다.
    draft = REPO / "docs" / "manuscript" / "methods_simulation_v7_draft.md"
    chk("draft 파일이 실재한다", draft.exists(), str(draft))
    if draft.exists():
        dtext = draft.read_text(encoding="utf-8")
        #  두 규약의 네 σ 값과 두 비 — 이것이 갈라지면 공저자가 서로 다른 표를 본다
        for key in ["54.0", "70.6", "72.3", "81.3", "1.308", "1.124"]:
            chk(f"draft 에도 있다: {key}", key in dtext)
        #  구조 지표: docx 에만 있고 draft 에 없으면 draft 가 낡은 것이다
        for key in ["86.6", "13.1", "15.5"]:
            chk(f"draft 에도 있다 (구조 지표): {key}", key in dtext)
        #  철회 세대 값이 **살아 있는 본문값**으로 새어 들어가지 않았는가.
        #  ⚠ 줄 단위로 본다 — draft 도 이 파일의 MAINTEXT_EDITS 처럼 "v6 원문 → 고침"
        #    대조 행에서 옛 값을 **의도적으로 인용**한다 (그것을 지우면 무엇을 고치는지
        #    알 수 없다).  그러므로 "옛 값이 있다" 가 아니라 **"옛 값이 있는데 같은 줄에
        #    새 값이 없다"** 가 결함이다.
        _stale = [ln for ln in dtext.splitlines()
                  if ("1.98 to 3.00" in ln or "1.98 → 3.00" in ln) and "54.0" not in ln]
        chk("draft 의 옛 세대 σ 는 고침 대조 행에서만 인용된다",
            not _stale, _stale[:1])

    #  9. σ_ion 보류 사유가 '완주 대기' 로 되돌아가지 않았는가.  값을 못 쓰는 이유는
    #     팔 수가 아니라 입력 규약이다 — 그 구분이 사라지면 독자가 기다리면 된다고 읽는다.
    chk("σ_ion 보류 사유가 팔 수 얘기가 아니다",
        "진행 중" not in PENDING_ION and "규약" in PENDING_ION, PENDING_ION)
    _ion_open = " ".join(c for k, c in OPEN_ITEMS if "σ_ion" in k)
    chk("미결 항목이 '남은 팔이 끝나도 안 바뀐다' 를 명시한다",
        "바뀌지 않습니다" in _ion_open)

    print(f"\n{len(fails)} failure(s)")
    return 1 if fails else 0


def main(argv=None):
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("-o", "--out", type=Path, default=DEFAULT_OUT)
    ap.add_argument("--selftest", action="store_true")
    args = ap.parse_args(argv)
    if args.selftest:
        return selftest()
    p = build(args.out)
    print(f"wrote {p}  ({p.stat().st_size:,} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
